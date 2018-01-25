import docx

doc = docx.Document('./data/CustomStatBlock.docx')

start_index = 0
beginning = []
for index in range(start_index, len(doc.paragraphs)):
    if "DEFENSE" in doc.paragraphs[index].text:
        start_index = index
        break
    else:
        beginning.append(doc.paragraphs[index])


def process_begin(stuff=None):
    class_list = [
        "bard", "cleric", "druid",
        "fighter", "wizard", "rogue",
        "ranger", "monk", "sorcerer"
    ]
    if not stuff:
        print("ERROR: nothing to process")
        pass
    returning = {}
    returning["name"] = stuff[0]
    temp = stuff[1].split()
    if temp[0].lower() in ["male", "female", "none"]:
        returning["name"] = temp[0]
    if temp[1].isaplha():
        returning["age"]
    return returning

defense = []
for index in range(start_index, len(doc.paragraphs)):
    if "OFFENSE" in doc.paragraphs[index].text:
        start_index = index
        break
    else:
        defense.append(doc.paragraphs[index])

offense = []
for index in range(start_index, len(doc.paragraphs)):
    if "STATISTICS" in doc.paragraphs[index].text:
        start_index = index
        break
    else:
        offense.append(doc.paragraphs[index])

statistics = []
for index in range(start_index, len(doc.paragraphs)):
    if "SPECIAL ABILITIES" in doc.paragraphs[index].text:
        start_index = index
        break
    else:
        statistics.append(doc.paragraphs[index])

special_abilities = []
for index in range(start_index, len(doc.paragraphs)):
    special_abilities.append(doc.paragraphs[index])

for item in beginning:
    print(item.text)
for item in defense:
    print(item.text)
for item in offense:
    print(item.text)
for item in statistics:
    print(item.text)
for item in special_abilities:
    print(item.text)

name = doc.paragraphs[0].text

temp = doc.paragraphs[1].text
line = temp.split(" ")
gender = line[0]
age = line[1]
subrace = line[2]  # check if exists
race = line[3]
classes = line[4]
"""
check for multiple of classes
if single class:
    class = line[4]
elif:
    class = line[4]
    multiclass = True
"""
# check for archetype
if line[5][0] == '(':
    archetype = line[5]
else:
    archetype = None

level = line[-1]
""" 
line = paragraph[1] split by space
gender = line[0]
age = line[1]
subrace = line[2] if exists
race = line[3]
check for multiple of classes
if single class:
    class = line[4]
elif:
    class = line[4]
    multiclass = True
archetype = line[5] if exists
level = line[-1] 
"""

temp = doc.paragraphs[2].text
line = temp.split(" ")
alignment = line[0]
size = line[1]
creature_type = line[2]
if len(line) > 3:
    subtype = line[3]
"""
line = paragraph[2] split by space
alignment = line[0]
size = line[1]
type = line[2]
if len(line) > 3:
    subtype = line[3]
"""

temp = doc.paragraphs[3].text
line = temp.split(';')
init_mod = line[0]
testing = "test"
if line[1].lower == "senses":
    senses = line[1]
perception = line[-1]
"""
line = paragraph[3] split by ';'
init_mod = line[0]
if senses:
    senses = line[1]
perception = line[-1]
"""

phrase = "DEFENSE"
index = 4
temp = []
attack = False
while index < len(doc.paragraphs):
    temp.append(doc.paragraphs[index].text)
    if not attack:
        if "DEFENSE" in temp[-1].upper():
            phrase = "ATTACK"
            attack = True
    elif "ATTACK" in temp[-1].upper():
        temp.pop()
        index = len(doc.paragraphs)
        break
    index += 1
lines = temp
for line in lines:
    if line[:1] == "AC":
        temp = line.split(",", 3)
        temp1 = temp[-1].split('(')
        temp[-1] = temp1
        AC = temp[0]
        touch = temp[1]
        temp2 = temp[2]
        flat_footed = temp[2]
        AC_breakdown = temp[-1]
    elif line[:1] == 'hp':
        temp = line.split(';')
        HP = temp[0]
        DR = temp[1]
    elif line[:3].lower() == 'fort':
        temp = line.split(',')
        fort = temp[0]
        reflex = temp[1]
        will = temp[2]
    elif line[:5].lower() == 'immune':
        temp = line.split(';')
        immune = temp[0]
        resist = temp[1]
        SR = temp[2]
    elif "defensive" in line[:9].lower():
        defensive_abilities = line
"""
!!!DEFENSE!!!
defenses = regex of between DEFENSE and OFFENSE
line = defenses[2] split by regex 
AC = line[0]
touch = line[1]
flat_footed = line[2]
AC_breakdown = line[3]

line = defenses[3] split by ';'
hp = line[0] split by regex
DR = line[1] split by regex
saves = defenses[4] split by regex
defensive_abilities = defenses[5] split by regex
"""

phrase = "str"
index = 8
temp = []
STR = False
while index < len(doc.paragraphs):
    temp.append(doc.paragraphs[index].text)
    if not STR:
        pass
    index += 1
    pass
print('finished')
"""
!!!OFFENSE!!!
offenses = regex between OFFENSE and stats
speed = offenses[2] split by regex 
melee = offenses[3] split by regex
ranged = offenses[4] split by regex
special_attacks = offenses[5:7]
spells = offenses[8:]

!!!STAT!!!
statistics = stat lines split by line
stats = statistics[0] split by ' '

line = statistics[1]
BAB = line[0]
CMB
CMD

feats = statistics[2] split by ','
skills = statistics[3] split by ','
if racial modifiers:
    racial_modifiers = statistics[4] split by regex
    add 1 to index to search for
traits = statistics[5] split by ','
languages = statistics[6] split by ','
SQ = statistics[7] split by regex
inventory = statistics[8] split by regex
    combat_gear, other_gear

!!!ABILITIES!!!
abilities = end of document split by regex
    extraordinary
    supernatural
    spell_like
"""
